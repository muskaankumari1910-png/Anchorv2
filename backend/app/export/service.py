from typing import List, Dict
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.models import Requirement, Evidence, Contradiction, GroundingStatus
import io


def export_to_docx(
    source_id: str,
    db: Session,
    workspace_id: str,
    include_quarantined: bool = False
) -> io.BytesIO:
    """
    Export requirements to .docx with traceability appendix.
    
    Sprint 8: Now scoped to workspace.
    Returns BytesIO buffer containing the document.
    """
    # Fetch requirements
    query = db.query(Requirement).join(Evidence).filter(
        Evidence.source_id == source_id,
        Evidence.workspace_id == workspace_id,
        Requirement.workspace_id == workspace_id,
        Requirement.is_merged == 0
    ).distinct()
    
    if not include_quarantined:
        query = query.filter(Requirement.grounding == GroundingStatus.GROUNDED)
    
    requirements = query.all()
    
    # Check for unresolved conflicts
    conflicts = db.query(Contradiction).join(
        Requirement, Contradiction.requirement_id_1 == Requirement.id
    ).join(
        Evidence, Evidence.requirement_id == Requirement.id
    ).filter(
        Evidence.source_id == source_id,
        Contradiction.status == "open"
    ).distinct().all()
    
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading('Requirements Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Warning about conflicts
    if conflicts:
        warning = doc.add_paragraph()
        warning.add_run('⚠️ WARNING: ').bold = True
        warning.add_run(f'{len(conflicts)} unresolved conflicts exist. Review before finalizing.')
        warning.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_paragraph()
    
    # Group by category
    by_category: Dict[str, List[Requirement]] = {}
    for req in requirements:
        cat = req.category or "Uncategorized"
        if cat not in by_category:
            by_category[cat] = []
        by_category[cat].append(req)
    
    # Write requirements by category
    for category, reqs in sorted(by_category.items()):
        doc.add_heading(category, level=1)
        
        for req in reqs:
            # Requirement statement
            p = doc.add_paragraph()
            p.add_run(f'• {req.statement}').bold = True
            
            # Metadata
            meta = doc.add_paragraph(style='List Bullet 2')
            meta.add_run(f'Type: {req.type.value} | ')
            meta.add_run(f'Confidence: {req.confidence.value} | ')
            meta.add_run(f'ID: {req.id}')
            meta.runs[0].font.size = Pt(9)
            meta.runs[0].font.color.rgb = RGBColor(100, 100, 100)
            
            doc.add_paragraph()
    
    # Traceability Appendix
    doc.add_page_break()
    doc.add_heading('Traceability Appendix', level=1)
    
    p = doc.add_paragraph()
    p.add_run('This appendix maps each requirement to its source quotes, enabling verification.')
    p.runs[0].font.italic = True
    
    doc.add_paragraph()
    
    for req in requirements:
        # Requirement ID and statement
        doc.add_heading(f'{req.id}', level=2)
        doc.add_paragraph(req.statement)
        
        # Evidence
        evidence_list = db.query(Evidence).filter(
            Evidence.requirement_id == req.id,
            Evidence.verified == 1
        ).all()
        
        if evidence_list:
            doc.add_paragraph('Evidence:', style='Heading 3')
            for idx, evd in enumerate(evidence_list, 1):
                evd_p = doc.add_paragraph(style='List Number')
                evd_p.add_run(f'Source: {evd.source_id}, Segment: {evd.segment_id}\n')
                evd_p.add_run(f'Quote: "{evd.verbatim_quote}"')
                evd_p.runs[1].font.italic = True
        
        doc.add_paragraph()
    
    # Statistics
    doc.add_page_break()
    doc.add_heading('Document Statistics', level=1)
    
    total_reqs = len(requirements)
    grounded = len([r for r in requirements if r.grounding == GroundingStatus.GROUNDED])
    
    doc.add_paragraph(f'Total Requirements: {total_reqs}')
    doc.add_paragraph(f'Grounded Requirements: {grounded}')
    doc.add_paragraph(f'Unresolved Conflicts: {len(conflicts)}')
    
    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def export_to_markdown(
    source_id: str,
    db: Session,
    workspace_id: str,
    include_quarantined: bool = False
) -> str:
    """
    Export requirements to markdown with traceability appendix.
    
    Sprint 8: Now scoped to workspace.
    Returns markdown string.
    """
    # Fetch requirements
    query = db.query(Requirement).join(Evidence).filter(
        Evidence.source_id == source_id,
        Evidence.workspace_id == workspace_id,
        Requirement.workspace_id == workspace_id,
        Requirement.is_merged == 0
    ).distinct()
    
    if not include_quarantined:
        query = query.filter(Requirement.grounding == GroundingStatus.GROUNDED)
    
    requirements = query.all()
    
    # Check conflicts
    conflicts = db.query(Contradiction).join(
        Requirement, Contradiction.requirement_id_1 == Requirement.id
    ).join(
        Evidence, Evidence.requirement_id == Requirement.id
    ).filter(
        Evidence.source_id == source_id,
        Contradiction.status == "open"
    ).distinct().all()
    
    lines = []
    
    # Title
    lines.append('# Requirements Document\n')
    
    # Warning
    if conflicts:
        lines.append(f'> ⚠️ **WARNING**: {len(conflicts)} unresolved conflicts exist. Review before finalizing.\n')
    
    lines.append('')
    
    # Group by category
    by_category: Dict[str, List[Requirement]] = {}
    for req in requirements:
        cat = req.category or "Uncategorized"
        if cat not in by_category:
            by_category[cat] = []
        by_category[cat].append(req)
    
    # Requirements by category
    for category, reqs in sorted(by_category.items()):
        lines.append(f'## {category}\n')
        
        for req in reqs:
            lines.append(f'- **{req.statement}**')
            lines.append(f'  - Type: {req.type.value} | Confidence: {req.confidence.value} | ID: `{req.id}`')
            lines.append('')
    
    # Traceability Appendix
    lines.append('---\n')
    lines.append('# Traceability Appendix\n')
    lines.append('*This appendix maps each requirement to its source quotes.*\n')
    
    for req in requirements:
        lines.append(f'### {req.id}\n')
        lines.append(f'{req.statement}\n')
        lines.append('**Evidence:**\n')
        
        evidence_list = db.query(Evidence).filter(
            Evidence.requirement_id == req.id,
            Evidence.verified == 1
        ).all()
        
        for idx, evd in enumerate(evidence_list, 1):
            lines.append(f'{idx}. Source: `{evd.source_id}`, Segment: `{evd.segment_id}`')
            lines.append(f'   > "{evd.verbatim_quote}"')
            lines.append('')
        
        lines.append('')
    
    # Statistics
    lines.append('---\n')
    lines.append('# Document Statistics\n')
    
    total_reqs = len(requirements)
    grounded = len([r for r in requirements if r.grounding == GroundingStatus.GROUNDED])
    
    lines.append(f'- **Total Requirements**: {total_reqs}')
    lines.append(f'- **Grounded Requirements**: {grounded}')
    lines.append(f'- **Unresolved Conflicts**: {len(conflicts)}')
    
    return '\n'.join(lines)
